def convertBaseball(self):
        # Reading docx file
        import sys
        from docx import Document
        document = Document(oldDocName)
        tables = document.tables
        section = []
        i = 0
        for table in tables:
            section.append(table)
        # document.save('MEN Game Notes Hitting_Updated.docx')
        #
        working = section[1]
        # self.write("Reading document...\n")
        content = []
        flag = 0
        lastNameCounter = 0
        # Import library for reading website
        from bs4 import BeautifulSoup
        import urllib2
       # import pandas as pd
        # download html from link
        url = "http://www.dakstats.com/WebSync/Pages/Team/IndividualStats.aspx?association=10&sg=MBA&sea=NAIMBA_2018&team=9101"
        page = urllib2.urlopen(url)
        soup = BeautifulSoup(page, "html.parser")
        table = soup.find("table", {"class": "gridViewReportBuilderWide"})
        avg_scores = []
        names = []
        count = 1
        dic = {}
        row_counter = 0
        for row in table.find_all('tr')[1:]:
            col = row.find_all('td')
            row_counter = row_counter + 1
            if len(col) == 22:
                name = col[0].find(text=True)
                names.append(name)
                GP = col[1].find(text=True)
                #
                # average = field 1
                average = col[3].find(text=True)
                # OB = field 2
                OB = col[17].find(text=True)
                # SLG = field 3
                SLG = col[12].find(text=True)
                #
                avg_scores.append(average)
                R = col[5].find(text=True)
                # 2B
                TwoB = col[7].find(text=True)
                # 3B
                ThreeB = col[8].find(text=True)
                # HR
                HR = col[9].find(text=True)
                # BB
                BB = col[13].find(text=True)
                # K not found
                # SB
                SB = col[20].find(text=True)

                try:
                    for row in working.rows:
                        for cell in row.cells:
                            for index, paragraph in enumerate(cell.paragraphs):
                                content.append(paragraph.text)
                                switch = False
                                if flag == 1:
                                    # print "doc: " + paragraph.text
                                    # print "website: " + average
                                    paragraph.text = average
                                    flag = flag + 1
                                    continue
                                if flag == 2:
                                    # print "doc: " + paragraph.text
                                    # print "website: " + OB
                                    paragraph.text = OB
                                    flag = flag + 1
                                    continue
                                if flag == 3:
                                    # print "doc: " + paragraph.text
                                    # print "website: " + SLG
                                    paragraph.text = SLG
                                    flag = flag + 1
                                    continue
                                if flag > 3:
                                    # print paragraph.text + '\n'


                                    if paragraph.text.__contains__('2B'):
                                        print '2B found'
                                        paragraph.text = "2B - " + TwoB

                                    if paragraph.text.__contains__('3B'):
                                        print '3B found'
                                        paragraph.text = "3B - " + ThreeB

                                    if paragraph.text.__contains__('HR'):
                                        print 'HR found'
                                        paragraph.text = "HR - " + HR

                                    if paragraph.text.__contains__('BB'):
                                        print 'BB found'
                                        paragraph.text = "BB - " + BB

                                    if paragraph.text.__contains__('K'):
                                        print 'K found'

                                    if paragraph.text.__contains__('SB'):
                                        print 'SB found'
                                        paragraph.text = "SB - " + SB

                                    if paragraph.text == (name.split(', ')[0]):
                                        flag = 0
                                        break

                                if paragraph.text == name.split(' ')[1]:
                                    flag = flag + 1
                                    print paragraph.text
                                    # scan has a dictionary of all the names found on the website
                                    # the program will check whether or not a first name has already been recorded
                                    # if it has, then the counter will go up from 1 to 2
                                    # if the counter is more than 1, then the program will have to match the specific name
                                    # from the document with a matching number
                                    scan = dic.keys()
                                    for index in scan:
                                        if index.__contains__(name.split(' ')[1]):
                                            switch = True
                                            count = count + 1
                                            dic[name] = count
                                            count = 1
                                    if switch == True:
                                        continue
                                    dic[name] = count

                except IndexError:
                    print "list index out of range"

        print "almost done"
        print "Check on these players' stats."
        for name in dic.keys():
            if dic[name] > 1:
                print name + '\n'
        print " Information might be wrong due to multiple first names appearing."
        document.add_table(1, 2, style=None)
        document.save(newDocDirectory + "/Baseball Updated File.docx")
        sys.exit()