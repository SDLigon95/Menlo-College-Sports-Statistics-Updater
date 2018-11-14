#!/usr/bin/env python2.7
# -*- coding: utf8 -*-
from Tkinter import *
from tkFileDialog import *
import os
# encoding=utf8
import sys

reload(sys)
sys.setdefaultencoding('utf8')
win = Tk()
win.geometry('575x357')  # Size 200, 200
choice = 0
menu_child = 0 
menu_child_switch = 0


class UI(Frame):
    oldDocName = []
    newDocDirectory = []

    def constructButtons(self, choice):
        self.openButton = Button(win, text="Open Word Document", command=self.openFile)
        self.img_openButton = PhotoImage(file="./open_doc.gif")
        self.openButton.config(image=self.img_openButton)
        self.openButton.pack()
        self.saveAsButton = Button(win, text="Save As", command=self.saveFileAs)
        self.img_saveAsButton = PhotoImage(file="./save_as.gif")
        self.saveAsButton.config(image=self.img_saveAsButton)
        self.saveAsButton.pack()
        if (choice == 1):
            self.convertButton = Button(win, text="Convert Baseball Stats", command=self.convertBaseball)
        if (choice == 2):
            self.convertButton = Button(win, text="Convert Baseball Stats", command=self.convertBaseballPitching)

        if (choice == 3):
            self.convertButton = Button(win, text="Convert Women's Soccer Stats", command=self.convertWomanSoccer)
        if (choice == 4):
            self.convertButton = Button(win, text="Convert Men's Soccer Stats", command=self.convertMenSoccer)
        if (choice == 5):
            self.convertButton = Button(win, text="Convert Volleyball Stats", command=self.convertVolleyball)
        if (choice == 6):
            self.convertButton = Button(win, text="Convert Women's Basketball Stats", command=self.convertWomanBasketball)
        if (choice == 7):
            self.convertButton = Button(win, text="Convert Men's Basketball Stats", command=self.convertManBasketball)
        if (choice == 8):
            self.convertButton = Button(win, text="Convert Softball Stats", command=self.convertSoftball)
        self.img_convertButton = PhotoImage(file="./convert.gif")
        self.convertButton.config(image=self.img_convertButton)
        self.convertButton.pack()

    def destroyButtons(self, menu_child):
        
        self.openButton.destroy()
        self.saveAsButton.destroy()
        self.convertButton.destroy()
        self.backButton.destroy()
        if (menu_child == 1):
            self.baseballMenu()
        if (menu_child == 2):
            self.soccerMenu()
        if (menu_child == 3):
            self.basketballMenu()
        if (menu_child == 0):
            self.create_UI()
        print str(menu_child)
        
      

    # Program startup
    def __init__(self):
        self.create_UI()

    #Create Main Menu UI    
    def create_UI(self):
        Frame.__init__(self)
        self.baseballOption = Button(win, text= "Baseball", command=self.baseballMenu)
        _file = os.path.abspath(sys.argv[0])
        path = os.path.dirname(_file)
        image_path = os.path.join(path, 'baseball.gif')
        self.img_baseball = PhotoImage(file=image_path)
        self.baseballOption.config(image=self.img_baseball)
        self.baseballOption.pack()
        self.softballOption = Button(win, text="Softball", command = self.softballMenu)
        _file = os.path.abspath(sys.argv[0])
        path = os.path.dirname(_file)
        image_path = os.path.join(path, 'softball.gif')
        print str(image_path)
        self.img_softball = PhotoImage(file=image_path)
        self.softballOption.config(image=self.img_softball)
        self.softballOption.pack()
        self.volleyballOption = Button(win, text="Volleyball", command=self.volleyballMenu)
        _file = os.path.abspath(sys.argv[0])
        path = os.path.dirname(_file)
        image_path = os.path.join(path, 'volleyball.gif')
        self.img_volleyball = PhotoImage(file=image_path)
        self.volleyballOption.config(image=self.img_volleyball)
        self.volleyballOption.pack()
        self.soccerOption = Button(win, text="Soccer", command=self.soccerMenu)
        _file = os.path.abspath(sys.argv[0])
        path = os.path.dirname(_file)
        image_path = os.path.join(path, 'soccer.gif')
        self.img_soccer = PhotoImage(file=image_path)
        self.soccerOption.config(image=self.img_soccer)
        self.soccerOption.pack()
        self.basketballOption = Button(win, text="Basketball", command=self.basketballMenu)
        _file = os.path.abspath(sys.argv[0])
        path = os.path.dirname(_file)
        image_path = os.path.join(path, 'basketball.gif')
        self.img_basketball = PhotoImage(file=image_path)
        self.basketballOption.config(image=self.img_basketball)
        self.basketballOption.pack()


    def destroyMenuChild(self,menu_child_switch):
        print "This is used"
        if (menu_child_switch ==1):
            self.baseballHittersOption.destroy()
            self.baseballPitchersOption.destroy()
            
        if menu_child_switch ==2:
            self.womanSoccerOption.destroy()
            self.menSoccerOption.destroy()
          
        if menu_child_switch ==3:
            self.womanBasketballOption.destroy()
            self.manBasketballOption.destroy()
        self.backButton.destroy()
        menu_child_switch = 0
        self.create_UI()
    #Used to destroy main menu layout (used in submenus)
    def destroyMenuUI(self):
        self.baseballOption.destroy()
        self.softballOption.destroy()
        self.volleyballOption.destroy()
        self.soccerOption.destroy()
        self.basketballOption.destroy()

    # def returnMenuChild(self,menu_child):
    #     if (menu_child==1):
            
    def printTest(self):
        print "TESTING"
  


    def baseballMenu(self):
        menu_child = 0
        menu_child_switch = 1
        print ("You're here")
        Frame.__init__(self)
        self.destroyMenuUI()
        self.baseballHittersOption = Button(win, text="Baseball Hitters", command=self.baseballHitterMenu)
        self.img_baseballHitter = PhotoImage(file="./baseball_hitter.gif")
        self.baseballHittersOption.config(image=self.img_baseballHitter)
        self.baseballHittersOption.pack()
        self.baseballPitchersOption = Button(win, text="Baseball Pitchers", command=self.baseballPitchingMenu)
        self.img_baseballPitcher = PhotoImage(file="./baseball_pitcher.gif")
        self.baseballPitchersOption.config(image=self.img_baseballPitcher)
        self.baseballPitchersOption.pack()
        self.backButton = Button(win, text="Back", command= lambda: self.destroyMenuChild(menu_child_switch))
        self.img_backButton = PhotoImage(file="./back_button.gif")
        self.backButton.config(image=self.img_backButton)
        self.backButton.pack(side=LEFT)
        # self.pack()
        

    def soccerMenu(self):
        menu_child == 0
        menu_child_switch = 2
        print ("You're here")
        Frame.__init__(self)
        self.destroyMenuUI()
        self.womanSoccerOption = Button(win, text="Women's Soccer", command=self.womanSoccerMenu)
        self.img_womenSoccer = PhotoImage(file="./women_soccer.gif")
        self.womanSoccerOption.config(image=self.img_womenSoccer)
        self.womanSoccerOption.pack()
        self.menSoccerOption = Button(win, text="Men's Soccer", command=self.menSoccerMenu)
        self.img_menSoccer = PhotoImage(file="./men_soccer.gif")
        self.menSoccerOption.config(image=self.img_menSoccer)
        self.menSoccerOption.pack()
        self.backButton = Button(win, text="Back", command= lambda: self.destroyMenuChild(menu_child_switch))
        self.img_backButton = PhotoImage(file="./back_button.gif")
        self.backButton.config(image=self.img_backButton)
        self.backButton.pack(side=LEFT)

    def basketballMenu(self):
        menu_child == 0
        menu_child_switch = 3
        print ("You're here")
        Frame.__init__(self)
        self.destroyMenuUI()
        self.womanBasketballOption = Button (win, text="Women's Basketball", command=self.womanBasketballMenu)
        self.img_womenBasketball = PhotoImage(file="./women_basketball.gif")
        self.womanBasketballOption.config(image=self.img_womenBasketball)
        self.womanBasketballOption.pack()
        self.manBasketballOption = Button(win, text="Men's Basketball", command=self.manBasketballMenu)
        self.img_menBasketball = PhotoImage(file="./men_basketball.gif")
        self.manBasketballOption.config(image=self.img_menBasketball)
        self.manBasketballOption.pack()
        self.backButton = Button(win, text="Back", command= lambda: self.destroyMenuChild(menu_child_switch))
        self.img_backButton = PhotoImage(file="./back_button.gif")
        self.backButton.config(image=self.img_backButton)
        self.backButton.pack(side=LEFT)


    def baseballHitterMenu(self):
        menu_child = 1
        self.backButton.destroy()
        self.baseballPitchersOption.destroy()
        self.baseballHittersOption.destroy()
        self.backButton = Button(win, text="Back", command=lambda: self.destroyButtons(menu_child))
        self.img_backButton = PhotoImage(file="./back_button.gif")
        self.backButton.config(image=self.img_backButton)
        self.backButton.pack(side=LEFT)
        choice = 1
        self.constructButtons(choice)

        # Background and Pack
        self.configure(background="white")
        self.pack()
    def baseballPitchingMenu(self):
        menu_child = 1
        self.backButton.destroy()
        self.baseballPitchersOption.destroy()
        self.baseballHittersOption.destroy()
        self.backButton = Button(win, text="Back", command=lambda: self.destroyButtons(menu_child))
        self.img_backButton = PhotoImage(file="./back_button.gif")
        self.backButton.config(image=self.img_backButton)
        self.backButton.pack(side=LEFT)
        choice = 2
        self.constructButtons(choice)

        # Background and Pack
        self.configure(background="white")
        self.pack()

    def softballMenu(self):
        menu_child = 0 
        self.volleyballOption.destroy()
        self.baseballOption.destroy()
        self.softballOption.destroy()
        self.soccerOption.destroy()
        self.basketballOption.destroy()
        self.backButton = Button(win, text="Back", command= lambda: self.destroyButtons(menu_child))
        self.img_backButton = PhotoImage(file="./back_button.gif")
        self.backButton.config(image=self.img_backButton)
        self.backButton.pack(side=LEFT)
        choice = 8
        self.constructButtons(choice)

        # Background and Pack
        self.configure(background="white")
        self.pack()

    def womanSoccerMenu(self):
        menu_child = 2
        self.backButton.destroy()
        self.womanSoccerOption.destroy()
        self.menSoccerOption.destroy()
        self.backButton = Button(win, text="Back", command=lambda: self.destroyButtons(menu_child))
        self.img_backButton = PhotoImage(file="./back_button.gif")
        self.backButton.config(image=self.img_backButton)
        self.backButton.pack(side=LEFT)
        choice = 3
        self.constructButtons(choice)

        # Background and Pack
        self.configure(background="white")
        self.pack()

    def womanBasketballMenu(self):
        menu_child = 3
        self.backButton.destroy()
        self.womanBasketballOption.destroy()
        self.manBasketballOption.destroy()
        self.backButton = Button(win, text="Back", command=lambda: self.destroyButtons(menu_child))
        self.img_backButton = PhotoImage(file="./back_button.gif")
        self.backButton.config(image=self.img_backButton)
        self.backButton.pack(side=LEFT)
        choice = 6
        self.constructButtons(choice)

        # Background and Pack
        self.configure(background="white")
        self.pack()

    def manBasketballMenu(self):
        menu_child = 3
        self.backButton.destroy()
        self.womanBasketballOption.destroy()
        self.manBasketballOption.destroy()
        self.backButton = Button(win, text="Back", command=lambda: self.destroyButtons(menu_child))
        self.img_backButton = PhotoImage(file="./back_button.gif")
        self.backButton.config(image=self.img_backButton)
        self.backButton.pack(side=LEFT)
        choice = 7
        self.constructButtons(choice)

        # Background and Pack
        self.configure(background="white")
        self.pack()

    def menSoccerMenu(self):
        menu_child = 2
        self.backButton.destroy()
        self.womanSoccerOption.destroy()
        self.menSoccerOption.destroy()
        self.backButton = Button(win, text="Back", command=lambda: self.destroyButtons(menu_child))
        self.img_backButton = PhotoImage(file="./back_button.gif")
        self.backButton.config(image=self.img_backButton)
        self.backButton.pack(side=LEFT)
        choice = 4
        self.constructButtons(choice)

        # Background and Pack
        self.configure(background="white")
        self.pack()


    def volleyballMenu(self):
        menu_child = 0 
        self.volleyballOption.destroy()
        self.baseballOption.destroy()
        # self.baseballOption.destroy()
        self.softballOption.destroy()
        self.soccerOption.destroy()
        self.basketballOption.destroy()
        self.backButton = Button(win, text="Back", command= lambda: self.destroyButtons(menu_child))
        self.img_backButton = PhotoImage(file="./back_button.gif")
        self.backButton.config(image=self.img_backButton)
        self.backButton.pack(side=LEFT)
        choice = 5
        self.constructButtons(choice)

        # Background and Pack
        self.configure(background="white")
        self.pack()

    # Grab file name and location from window
    def openFile(self):
        global oldDocName
        oldDocName = askopenfilename(filetypes=(("Word files", "*.docx"),))

        print(oldDocName)
        # return fileName

    def saveFileAs(self):
        global newDocDirectory
        newDocDirectory = askdirectory()
        if newDocDirectory is None:  # asksaveasfile return `None` if dialog closed with "cancel".
            return
        print newDocDirectory

        # Grabbing input to display in log window
        # def write (self, txt):
        #     self.output.insert(END,str(txt))
        #     self.update_idletasks()

    # Grabs input from website and compares names found in original document.
    # Scores will get updated based on which names are found.

    def convertSoftball(self):
            # Reading docx file
            import sys
            from docx import Document
            document = Document(oldDocName)
            print (oldDocName)
            tables = document.tables
            section = []
            i = 0
            for table in tables:
                section.append(table)
            # document.save('MEN Game Notes Hitting_Updated.docx')
            #
            working = section[1]
            # self.write("Reading document...\n")
            content = []
            flag = 0
            lastNameCounter = 0
            # Import library for reading website
            from bs4 import BeautifulSoup
            import urllib2
        # import pandas as pd
            # download html from link
            url = "http://www.dakstats.com/WebSync/Pages/Team/IndividualStats.aspx?association=10&sg=WBA&sea=NAIWBA_2018&team=1501"
            page = urllib2.urlopen(url)
            soup = BeautifulSoup(page, "html.parser")
            table = soup.find("table", {"class": "gridViewReportBuilderWide"})
            avg_scores = []
            names = []
            count = 0

            dic = {}
            row_counter = 0
            for row in table.find_all('tr')[1:]:
                col = row.find_all('td')
                row_counter = row_counter + 1
                if len(col) == 22:
                    name = col[0].find(text=True)
                    if name=="Kennedy, Jordan":
                        print "Kennedy, Jordan yay1234567890"
                        continue
                    if name=="Cervantes, Victoria":
                        print "Cervantes, Victoria yay1234567890"
                        continue
                    if name=="Martinez, Larissa":
                        print "Martinez, Larissa yay1234567890"
                        continue
                    names.append(name)
                    # if name.split(' ')[1] == "Victoria" or name.split(' ')[1] == "Jordan" or name.split(' ')[1] == "Larissa":
                    #     print ""
                    GP = col[1].find(text=True)
                    #
                    # average = field 1
                    average = col[3].find(text=True)
                    # OB = field 2
                    OB = col[17].find(text=True)
                    # SLG = field 3
                    SLG = col[12].find(text=True)
                    #
                    avg_scores.append(average)
                    R = col[5].find(text=True)
                    # 2B
                    TwoB = col[7].find(text=True)
                    # 3B
                    ThreeB = col[8].find(text=True)
                    # HR
                    HR = col[9].find(text=True)
                    #RBI
                    RBI = col[10].find(text=True)
                    # BB
                    BB = col[13].find(text=True)
                    # K not found
                    # SB
                    SO = col[15].find(text=True)
                    SB = col[20].find(text=True)

                    try:
                        for row in working.rows:
                            for cell in row.cells:
                                for index, paragraph in enumerate(cell.paragraphs):
                                    content.append(paragraph.text)
                                    switch = False
                                    # if paragraph.text == "Cassie" and name.split(' ')[1] == "Brooke":
                                    #     count = count + 1
                                    #     print "Cassie -> " + str(count)
                                            
                                        
                                    if flag == 1:
                                        print "1"
                                        # print "doc: " + paragraph.text
                                        # print "website: " + average
                                        paragraph.text = average
                                        flag = flag + 1
                                        continue
                                    if flag == 2:
                                        print "2"
                                        # print "doc: " + paragraph.text
                                        # print "website: " + OB
                                        paragraph.text = OB
                                        flag = flag + 1
                                        continue
                                    if flag == 3:
                                        print "3"
                                        # print "doc: " + paragraph.text
                                        # print "website: " + SLG
                                        paragraph.text = SLG
                                        flag = flag + 1
                                        continue
                                    if flag > 3:
                                        print "____*->" + paragraph.text + '\n'
                                        # print "FLAG==" + str(flag) + " "+ name
                                        if name.split(' ')[1].__contains__("Sarah"):
                                            print "This is something: " +  name.split(' ')[1]

                                        if paragraph.text.startswith("R-"):
                                            paragraph.text = "R-" + R
                                            print 'R found: '  + paragraph.text

                                        if paragraph.text.startswith("RBI-"):
                                            paragraph.text = "RBI-" + RBI
                                            print 'RBI found: ' + paragraph.text 

                                        if paragraph.text.startswith('2B-'):
                                            paragraph.text = "2B-" + TwoB
                                            print '2B found: ' + paragraph.text

                                        if paragraph.text.startswith('3B-'):
                                            paragraph.text = "3B-" + ThreeB
                                            print '3B found: ' + paragraph.text

                                        if paragraph.text.startswith('HR-'):
                                            paragraph.text = "HR- " + HR
                                            print 'HR found: ' + paragraph.text

                                        if paragraph.text.startswith('BB-'):
                                            paragraph.text = "BB- " + BB
                                            print 'BB found: ' + paragraph.text

                                        if paragraph.text.startswith('K-'):
                                            paragraph.text = "K-" + SO
                                            print 'K found: ' + paragraph.text

                                        if paragraph.text.startswith('SB-'):
                                            paragraph.text = "SB- " + SB
                                            print 'SB found' + paragraph.text
                                            flag = 0
                                            break

                                        # if paragraph.text == (name.split(', ')[0]):
                                        #     print "->END: " + paragraph.text
                                        #     flag = 0
                                        #     break

                                        # if paragraph.text != (name.split(', ')[0]):
                                        #     print "->NO_END: " + paragraph.text
                                        
                                    if paragraph.text == name.split(' ')[1]:
                                        if paragraph.text == "Brooke":
                                            count = count + 1
                                            print "LOOK OVER HERE" + str(count)
                                            
                                        # if paragraph.text  == "Victoria" or paragraph.text == "Jordan" or paragraph.text == "Larissa":
                                        #     print "-> INSERT NEW METHOD FOR " + paragraph.text
                                        flag = flag + 1
                                        print  "->START: " + paragraph.text 
                                        
                                        # scan has a dictionary of all the names found on the website
                                        # the program will check whether or not a first name has already been recorded
                                        # if it has, then the counter will go up from 1 to 2
                                        # if the counter is more than 1, then the program will have to match the specific name
                                        # from the document with a matching number
                                        # scan = dic.keys()
                                        # for index in scan:
                                        #     if index.__contains__(name.split(' ')[1]):
                                        #         switch = True
                                        #         count = count + 1
                                        #         dic[name] = count
                                        #         count = 1
                                        # if switch == True:
                                        #     continue
                                        # dic[name] = count

                    except IndexError:
                        print "list index out of range"

            print "almost done"
            print "Check on these players' stats."
            for name in dic.keys():
                if dic[name] > 1:
                    print name + '\n'
            print " Information might be wrong due to multiple first names appearing."
            # document.add_table(1, 2, style=None)
            document.save(newDocDirectory + "/Softball Updated File.docx")
            newDocName = "/Softball Updated File.docx"
            self.convertSoftball_LOB(newDocDirectory, newDocName)
            # sys.exit()

    def convertSoftball_LOB (self, newDocDirectory, newDocName):
            # Reading docx file
            import sys
            from docx import Document
            from docx.shared import Pt

            
            document = Document(newDocDirectory + newDocName)

            style = document.styles['Normal']
            font = style.font
            font.name = 'Arial Narrow'
            font.size = Pt(9)
            # font.bold = True

            tables = document.tables
            section = []
            i = 0
            for table in tables:
                section.append(table)
            # document.save('MEN Game Notes Hitting_Updated.docx')
            #
            working = section[1]
            # self.write("Reading document...\n")
            content = []
            flag = 0
            lastNameCounter = 0
            # Import library for reading website
            from bs4 import BeautifulSoup
            import urllib2
        # import pandas as pd
            # download html from link
            url = "http://www.dakstats.com/WebSync/Pages/Team/IndividualStats.aspx?association=10&sg=WBA&sea=NAIWBA_2018&team=1501"
            page = urllib2.urlopen(url)
            soup = BeautifulSoup(page, "html.parser")
            table = soup.find_all("table", {"class": "gridViewReportBuilderWide"})[1]
            avg_scores = []
            names = []
            count = 1
            dic = {}
            row_counter = 0
            for row in table.find_all('tr')[1:]:
                col = row.find_all('td')
                row_counter = row_counter + 1
                if len(col) == 26:
                    name = col[0].find(text=True)
                    if name=="Reyes, Sarah":
                        print "yay1234567890"
                        continue
                    names.append(name)
                    print name
                    #ERA-1 F2
                    #W-2 F1
                    #L-3 F1
                    #GP-4
                    # GS-5
                    # CG-6 F3
                    # IP-10
                    # H-11
                    # BB-14
                    # KS-??
                    # OBA-20
                    # HR-18

                    ERA = col[1].find(text=True)
                    W = col[2].find(text=True)
                    L = col[3].find(text=True)
                    GP = col[4].find(text=True)
                    GS = col[5].find(text=True)
                    CG = col[6].find(text=True)
                    IP = col[10].find(text=True)
                    H = col[11].find(text=True)
                    BB = col[14].find(text=True)
                    HR = col[18].find(text=True)
                    SO = col[15].find(text=True)
                    OBA = col[20].find(text=True)


                    try:
                        for row in working.rows:
                            for cell in row.cells:
                                for index, paragraph in enumerate(cell.paragraphs):
                                    content.append(paragraph.text)
                                    switch = False
                                    if flag == 1:
                                        # print "doc: " + paragraph.text
                                        # print "website: " + average
                                        paragraph.text = "W:"+W+"L:"+L
                                        flag = flag + 1
                                        continue
                                    if flag == 2:
                                        # print "doc: " + paragraph.text
                                        # print "website: " + OB
                                        paragraph.text = ERA
                                        flag = flag + 1
                                        continue
                                    if flag == 3:
                                        # print "doc: " + paragraph.text
                                        # print "website: " + SLG
                                        paragraph.text = CG
                                        flag = flag + 1
                                        continue
                                    if flag > 3:
                                        # print paragraph.text + '\n'

                                        if paragraph.text.startswith("GP-"):
                                            print 'GP found'
                                            paragraph.text = "GP-" + GP

                                        if paragraph.text.startswith("GS-"):
                                            print 'GS found'
                                            paragraph.text = "GS-" + GS

                                        if paragraph.text.startswith('IP-'):
                                            print 'IP found'
                                            paragraph.text = "IP-" + IP

                                        if paragraph.text.startswith('H-'):
                                            print 'H found'
                                            paragraph.text = "H-" + H

                                        if paragraph.text.startswith('BB-'):
                                            print 'BB found'
                                            paragraph.text = "BB-" + BB

                                        if paragraph.text.startswith('KS-'):
                                            print 'KS found'
                                            paragraph.text = "KS-" + SO
                                        
                                        if paragraph.text.startswith('OBA-'):
                                            print 'OBA found'
                                            paragraph.text = "OBA-" + OBA

                                        if paragraph.text.startswith('HR-'):
                                            print 'HR found'
                                            paragraph.text = "HR-" + HR
                                            flag = 0
                                            break

                        
                                            

                                    if paragraph.text == name.split(' ')[1]:
                                        # if paragraph.text  == "Victoria" or paragraph.text == "Jordan" or paragraph.text == "Larissa":
                                        #     print "-> INSERT NEW METHOD FOR " + paragraph.text
                                        flag = flag + 1
                                        print paragraph.text
                                        # scan has a dictionary of all the names found on the website
                                        # the program will check whether or not a first name has already been recorded
                                        # if it has, then the counter will go up from 1 to 2
                                        # if the counter is more than 1, then the program will have to match the specific name
                                        # from the document with a matching number
                                        scan = dic.keys()
                                        for index in scan:
                                            if index.__contains__(name.split(' ')[1]):
                                                switch = True
                                                count = count + 1
                                                dic[name] = count
                                                count = 1
                                        if switch == True:
                                            continue
                                        dic[name] = count

                    except IndexError:
                        print "list index out of range"

            print "almost done"
            print "Check on these players' stats."
            for name in dic.keys():
                if dic[name] > 1:
                    print name + '\n'
            print " Information might be wrong due to multiple first names appearing."
            # document.add_table(1, 2, style=None)
            document.save(newDocDirectory + newDocName)
            print "-> SUCCESS NEW METHOD"
            sys.exit()
    def convertBaseball(self):
        # Reading docx file
        import sys
        from docx import Document
        document = Document(oldDocName)
        tables = document.tables
        section = []
        i = 0
        for table in tables:
            section.append(table)
        # document.save('MEN Game Notes Hitting_Updated.docx')
        #
        working = section[1]
        # self.write("Reading document...\n")
        content = []
        flag = 0
        lastNameCounter = 0
        # Import library for reading website
        from bs4 import BeautifulSoup
        import urllib2
       # import pandas as pd
        # download html from link
        url = "http://www.dakstats.com/WebSync/Pages/Team/IndividualStats.aspx?association=10&sg=MBA&sea=NAIMBA_2018&team=9101"
        page = urllib2.urlopen(url)
        soup = BeautifulSoup(page, "html.parser")
        table = soup.find("table", {"class": "gridViewReportBuilderWide"})
        avg_scores = []
        names = []
        count = 1
        dic = {}
        row_counter = 0
        new_switch = FALSE
        for row in table.find_all('tr')[1:]:
            col = row.find_all('td')
            row_counter = row_counter + 1
            if len(col) == 22:
                name = col[0].find(text=True)
                if name == "Merken, Andrew":
                    break
                names.append(name)
                print "NAME: " + name
                GP = col[1].find(text=True)
                #
                # average = field 1
                average = col[3].find(text=True)
                print "AVG: " + average
                # OB = field 2
                OB = col[17].find(text=True)
                print "OB: " + OB
                # SLG = field 3
                SLG = col[12].find(text=True)
                print "SLG: " + SLG
                #
                avg_scores.append(average)
                R = col[5].find(text=True)
                print "R: " + R
                # 2B
                TwoB = col[7].find(text=True)
                print "TwoB: " + TwoB
                # 3B
                ThreeB = col[8].find(text=True)
                print "ThreeB: " + ThreeB
                # HR
                HR = col[9].find(text=True)
                print "HR: " + HR
                #RBI
                RBI = col[10].find(text=True)
                print "RBI: " + RBI
                # BB
                BB = col[13].find(text=True)
                print "BB: " + BB
                # SO/K/KS
                SO = col[15].find(text=True)
                print "SO: " + SO
                # K not found
                # SB
                SB = col[20].find(text=True)
                print "SB: " + SB

                try:
                    for row in working.rows:
                        for cell in row.cells:
                            for index, paragraph in enumerate(cell.paragraphs):
                                content.append(paragraph.text)
                                switch = False
                                if flag == 1:
                                    # print "doc: " + paragraph.text
                                    print "website: " + average
                                    paragraph.text = average
                                    flag = flag + 1
                                    continue
                                if flag == 2:
                                    # print "doc: " + paragraph.text
                                    print "website: " + OB
                                    paragraph.text = OB
                                    flag = flag + 1
                                    continue
                                if flag == 3:
                                    # print "doc: " + paragraph.text
                                    print "website: " + SLG
                                    paragraph.text = SLG
                                    flag = flag + 1
                                    continue
                                if flag > 3:
                                    # print paragraph.text + '\n'

                                    if paragraph.text == name.split(', ')[0]:
                                        print "----> " + paragraph.text + " PASSED!!!"
                                        new_switch = TRUE

                                    if paragraph.text.startswith('R-'):
                                        print 'R found ' + paragraph.text + str(R)+ " " + str(flag)
                                        paragraph.text = "R- " + R

                                    if paragraph.text.__contains__('RBI-'):
                                        print 'RBI found' + paragraph.text +  str(RBI) + " " + str(flag)
                                        paragraph.text = "RBI- " + RBI 

                                    if paragraph.text.__contains__('2B-'):
                                        print '2B found ' + paragraph.text +  str(TwoB)+ " " + str(flag)
                                        paragraph.text = "2B- " + TwoB

                                    if paragraph.text.__contains__('3B-'):
                                        print '3B found ' + paragraph.text +  str(ThreeB) + " " + str(flag)
                                        paragraph.text = "3B- " + ThreeB

                                    if paragraph.text.startswith('HR-'):
                                        print 'HR found' + paragraph.text +  str(HR) + " " + str(flag)
                                        paragraph.text = "HR- " + HR

                                    if paragraph.text.startswith('BB-'):
                                        print 'BB found' + paragraph.text +  str(BB)+ " " + str(flag)
                                        paragraph.text = "BB- " + BB

                                    if paragraph.text.__contains__('K-'):
                                        print 'K found' + paragraph.text +  str(SO)+ " " + str(flag)
                                        paragraph.text = "K- " + SO
                            
                                    if paragraph.text.__contains__('SB-'):
                                        print str(flag)
                                        print 'SB found' + paragraph.text +  str(SB)
                                        paragraph.text = "SB- " + SB
                                        flag = 0
                                        if new_switch != TRUE:
                                            print name.split(', ')[0]
                                            print "Mismatched Last Name DETECTED" 
                                        new_switch = FALSE
                                        
                                        break
                                    
                                    # else:
                                    #     print paragraph.text

                                    # if paragraph.text == (name.split(', ')[0]):
                                    #     print "-> END: " + name.split(', ')[0]
                                    #     flag = 0
                                    #     break
                                    
                                    # if paragraph.text != (name.split(', ')[1]):
                                    #     print "->NO_END: " + paragraph.text
                                    #     if paragraph.text == ("Alexander"):
                                    #         print "******************"
                                        

                                if paragraph.text == name.split(' ')[1]:
                                    flag = flag + 1
                                    print "-> BEGIN: " + paragraph.text
                                    # scan has a dictionary of all the names found on the website
                                    # the program will check whether or not a first name has already been recorded
                                    # if it has, then the counter will go up from 1 to 2
                                    # if the counter is more than 1, then the program will have to match the specific name
                                    # from the document with a matching number
                                    scan = dic.keys()
                                    for index in scan:
                                        if index.__contains__(name.split(' ')[1]):
                                            switch = True
                                            count = count + 1
                                            print "---> "+name.split(' ')[1] + " " + str(count)
                                            dic[name] = count
                                            # count = 1
                                            
                                    if switch == True:
                                        continue
                                    dic[name] = count

                except IndexError:
                    print "list index out of range"

        print "almost done"
        print "Check on these players' stats."
        for name in dic.keys():
            if dic[name] > 1:
                print name + '\n'
        print " Information might be wrong due to multiple first names appearing."
        document.add_table(1, 2, style=None)
        document.save(newDocDirectory + "/Baseball Updated File.docx")
        sys.exit()

    def convertBaseballPitching(self):
        # Reading docx file
        import sys
        from docx import Document 
        from docx.shared import Pt
        #
        document = Document(oldDocName)
        #
        style = document.styles['Normal']
        font = style.font
        font.name = 'Arial Narrow'
        font.size = Pt(9)
        #
        tables = document.tables
        section = []
        i = 0
        for table in tables:
            section.append(table)
        # document.save('MEN Game Notes Hitting_Updated.docx')
        #
        working = section[1]
        # self.write("Reading document...\n")
        content = []
        flag = 0
        lastNameCounter = 0
        # Import library for reading website
        from bs4 import BeautifulSoup
        import urllib2
       # import pandas as pd
        # download html from link
        url = "http://www.dakstats.com/WebSync/Pages/Team/IndividualStats.aspx?association=10&sg=MBA&sea=NAIMBA_2018&team=9101"
        page = urllib2.urlopen(url)
        soup = BeautifulSoup(page, "html.parser")
        table = soup.find_all("table", {"class": "gridViewReportBuilderWide"})[1]
        avg_scores = []
        names = []
        count = 1
        dic = {}
        row_counter = 0
        for row in table.find_all('tr')[1:]:
            col = row.find_all('td')
            row_counter = row_counter + 1
            if len(col) == 26:
                name = col[0].find(text=True)
                names.append(name)
                print "NAME: " + name                
                W = col[2].find(text=True)
                print "W: " + W
                L = col[3].find(text=True)
                print "L: " + L
                GP = col[4].find(text=True)
                print "GP: " + GP
                BatAVG = col[20].find(text=True)
                print "BatAVG: " + BatAVG
                ERA = col[1].find(text=True)
                print "ERA: " + ERA
                IP = col[10].find(text=True)
                print "IP: " + IP
                BB = col[14].find(text=True)
                print "BB: " + BB
                SO = col[15].find(text=True)
                print "SO: " + SO
                HBP = col[22].find(text=True)
                print "HBP: " + HBP
                #GP
                #BAVG
     

                try:
                    for row in working.rows:
                        for cell in row.cells:
                            for index, paragraph in enumerate(cell.paragraphs):
                                content.append(paragraph.text)
                                switch = False
                                if flag == 1:
                                    print "W-" + W
                                    print "L-" + L
                                    paragraph.text = "W-" + W + "L-" + L
                                    flag = flag + 1
                                    continue
                                if flag == 2:
                                    print "ERA-" +ERA
                                    paragraph.text = ERA
                                    flag = flag + 1
                                    continue
                                if flag == 3:
                                    # print "doc: " + paragraph.text
                                    print "IP-" + IP
                                    paragraph.text = IP
                                    flag = flag + 1
                                    continue
                                if flag > 3:
                                    if paragraph.text.__contains__('WHIP'):
                                        print "WHIP found"

                                    if paragraph.text.__contains__('K-'):
                                        print 'K found' + paragraph.text +  str(SO)+ " " + str(flag)
                                        paragraph.style = document.styles['Normal']
                                        paragraph.text = "K-" + SO

                                    if paragraph.text.startswith('BB'):
                                        print 'BB found ' + paragraph.text +  str(BB) + " " + str(flag)
                                        paragraph.text = "BB- " + BB

                                    if paragraph.text.__contains__('HBP'):
                                        print 'HBP found' + paragraph.text +  str(HBP) + " " + str(HBP)
                                        paragraph.text = "HBP- " + HBP

                                    if paragraph.text.__contains__('GP-'):
                                        print 'GP- found ' + paragraph.text +  str(GP) + " " + str(flag)
                                        paragraph.text = "GP- " + GP
                                    
                                    if paragraph.text.__contains__('B/Avg'):
                                        print 'B/Avg- found ' + paragraph.text +  str(BatAVG) + " " + str(flag)
                                        paragraph.text = BatAVG
                                        flag = 0
                                        break

                                    # if paragraph.text == (name.split(', ')[0]):
                                    #     print "-> END: " + name.split(', ')[0]
                                    #     flag = 0
                                    #     break
                                    
                                    # if paragraph.text != (name.split(', ')[1]):
                                    #     print "->NO_END: " + paragraph.text
                                    #     if paragraph.text == ("Alexander"):
                                    #         print "******************"
                                        
                                
                                if paragraph.text == name.split(' ')[1]:
                                    flag = flag + 1
                                    print "-> BEGIN: " + paragraph.text
                                    # scan has a dictionary of all the names found on the website
                                    # the program will check whether or not a first name has already been recorded
                                    # if it has, then the counter will go up from 1 to 2
                                    # if the counter is more than 1, then the program will have to match the specific name
                                    # from the document with a matching number
                                    scan = dic.keys()
                                    for index in scan:
                                        if index.__contains__(name.split(' ')[1]):
                                            switch = True
                                            count = count + 1
                                            dic[name] = count
                                            count = 1
                                    if switch == True:
                                        continue
                                    dic[name] = count

                except IndexError:
                    print "list index out of range"

        print "almost done"
        print "Check on these players' stats."
        for name in dic.keys():
            if dic[name] > 1:
                print name + '\n'
        print " Information might be wrong due to multiple first names appearing."
        document.add_table(1, 2, style=None)
        document.save(newDocDirectory + "/Baseball_Pitchers Updated File.docx")
        sys.exit()

    def convertGoalie(self, newDocName, gender):
        print "GOALIE BEGIN"
        print gender
        endGoalieCounter = 0
        url = ""
        # Reading docx file
        import sys
        from docx import Document
        document = Document(newDocName)
        tables = document.tables
        section = []
        i = 0
        for table in tables:
            section.append(table)
        working = section[1]
        content = []
        flag = 0
        lastNameCounter = 0
        # Import library for reading website
        from bs4 import BeautifulSoup
        import urllib2
        #import pandas as pd
        # download html from link
        #2017 link http://www.dakstats.com/WebSync/Pages/Team/IndividualStats.aspx?association=10&sg=WSO&sea=NAIWSO_2017&team=2409
        if gender == "woman":
            url = "http://www.dakstats.com/WebSync/Pages/Team/IndividualStats.aspx?association=10&sg=WSO&team=2409&sea=NAIWSO_2018"
        if gender == "man":
            url = "http://www.dakstats.com/WebSync/Pages/Team/IndividualStats.aspx?association=10&sg=MSO&team=2623&sea=NAIMSO_2018"
        page = urllib2.urlopen(url)
        soup = BeautifulSoup(page, "html.parser")
        table = soup.find_all("table", {"class": "gridViewReportBuilderWide"})[1]
        avg_scores = []
        names = []
        count = 1
        dic = {}
        row_counter = 0
        for row in table.find_all('tr')[1:]:
            col = row.find_all('td')
            row_counter = row_counter + 1
            # if row_counter > 14:
            #     break
            if len(col) == 13:
                print "SUCCEED13!!!"
                try:
                    name = col[1].find('a', href=True)
                    name = name.get('title')
                    names.append(name.split(', ')[1])
                    if name.split(', ')[1] == "Paula":
                        print "This is it"
                        print name.split(', ')[1]

                        GP = col[2].find(text=True)
                        print "GP: " + GP

                        #
                        # average = field 1
                        GS = col[3].find(text=True)
                        print "GS: " + GS

                        Min = col[4].find(text=True)
                        print "Min: " + Min

                        GA = col[5].find(text=True)
                        print "GA: " + GA

                        GAAvg = col[6].find(text=True)
                        print "GAAvg: " + GAAvg

                        SV = col[7].find(text=True)
                        print "SV: " + SV

                        Pct = col[8].find(text=True)
                        print "Pct: " + Pct

                        W = col[9].find(text=True)
                        print "W: " + W

                        L = col[10].find(text=True)
                        print "L: " + L

                        T = col[11].find(text=True)
                        print "T: " + T

                        SHO = col[12].find(text=True)
                        print "SHO: " + SHO

                        try:
                            print "TESTING"
                            for row in working.rows:
                                for cell in row.cells:
                                    for index, paragraph in enumerate(cell.paragraphs):
                                        content.append(paragraph.text)
                                        switch = False
                                    print str(flag)
                                    print paragraph.text
                                    print name.split(', ')[1]
                                    if flag == 1:
                                        print "W FOUND: " + W
                                        print "W: " + paragraph.text
                                        paragraph.text = W
                                        flag = flag + 1
                                        continue
                                    if flag == 2:
                                        print "L FOUND: " + L
                                        print "L: " + paragraph.text
                                        paragraph.text = L
                                        flag = flag + 1
                                        continue
                                    if flag == 3:
                                        print "T FOUND: " + T
                                        print "T: " + paragraph.text
                                        paragraph.text = T
                                        flag = flag + 1
                                        continue
                                    if flag > 3:
                                        if paragraph.text.__contains__('GP-'):
                                            print 'GP found: ' + GP
                                            paragraph.text = "GP-" + GP

                                        if paragraph.text.__contains__('GS-'):
                                            print 'GS found: ' + GS
                                            paragraph.text = "GS-" + GS

                                        if paragraph.text.__contains__('Min-'):
                                            print 'Min found: ' + Min
                                            paragraph.text = "Min-" + Min

                                        if paragraph.text.__contains__('GA-'):
                                            print 'GA found: ' + GA
                                            paragraph.text = "GA-" + GA

                                        if paragraph.text.__contains__('GAAvg-'):
                                            print 'GAAvg found: ' + GAAvg
                                            paragraph.text = "GAAvg-" + GAAvg

                                        if paragraph.text.__contains__('SV-'):
                                            print 'SV found: ' + SV
                                            paragraph.text = "SV-" + SV

                                        if paragraph.text.__contains__('Pct-'):
                                            print 'Pct found: ' + Pct
                                            paragraph.text = "Pct-" + Pct

                                        if paragraph.text.__contains__('SHO-'):
                                            print 'SHO found: ' + SHO
                                            paragraph.text = "SHO-" + SHO
                                            endGoalieCounter = endGoalieCounter + 1
                                            if endGoalieCounter == 2:
                                                print "END OF GOAL"
                                                document.save(newDocDirectory + "/Updated Filev2.docx")
                                                sys.exit()
                                            flag = 0
                                            print "ROW COUNTER: " + str(row_counter)
                                            break

                                    if paragraph.text.__contains__(name.split(', ')[1]):
                                        print "NAMES2: " + paragraph.text + " " + name.split(', ')[1]
                                        flag = flag + 1
                                        print paragraph.text
                                        print "TEST2"

                                        # scan has a dictionary of all the names found on the website
                                        # the program will check whether or not a first name has already been recorded
                                        # if it has, then the counter will go up from 1 to 2
                                        # if the counter is more than 1, then the program will have to match the specific name
                                        # from the document with a matching number
                                        scan = dic.keys()
                                        for index in scan:
                                            if index.__contains__(name.split(', ')[1]):
                                                switch = True
                                                count = count + 1
                                                dic[name] = count
                                                count = 1

                                        if switch == True:
                                            continue
                                        dic[name] = count
                                        # paragraph.text[index + 1] = '.555'
                                        # print paragraph.text[index + 1]
                            try:
                                workingSecondPage = section[3]
                                for row in workingSecondPage.rows:
                                    for cell in row.cells:
                                        for index, paragraph in enumerate(cell.paragraphs):
                                            content.append(paragraph.text)
                                            switch = False
                                        print paragraph.text
                                        print name.split(', ')[1]
                                        if flag == 1:
                                            print "W FOUND: " + W
                                            print "W: " + paragraph.text
                                            paragraph.text = W
                                            flag = flag + 1
                                            continue
                                        if flag == 2:
                                            print "L FOUND: " + L
                                            print "L: " + paragraph.text
                                            paragraph.text = L
                                            flag = flag + 1
                                            continue
                                        if flag == 3:
                                            print "T FOUND: " + T
                                            print "T: " + paragraph.text
                                            paragraph.text = T
                                            flag = flag + 1
                                            continue
                                        if flag > 3:
                                            if paragraph.text.__contains__('GP-'):
                                                print 'GP found: ' + GP
                                                paragraph.text = "GP-" + GP

                                            if paragraph.text.__contains__('GS-'):
                                                print 'GS found: ' + GS
                                                paragraph.text = "GS-" + GS

                                            if paragraph.text.__contains__('Min-'):
                                                print 'Min found: ' + Min
                                                paragraph.text = "Min-" + Min

                                            if paragraph.text.__contains__('GA-'):
                                                print 'GA found: ' + GA
                                                paragraph.text = "GA-" + GA

                                            if paragraph.text.__contains__('GAAvg-'):
                                                print 'GAAvg found: ' + GAAvg
                                                paragraph.text = "GAAvg-" + GAAvg

                                            if paragraph.text.__contains__('SV-'):
                                                print 'SV found: ' + SV
                                                paragraph.text = "SV-" + SV

                                            if paragraph.text.__contains__('Pct-'):
                                                print 'Pct found: ' + Pct
                                                paragraph.text = "Pct-" + Pct

                                            if paragraph.text.__contains__('SHO-'):
                                                print 'SHO found: ' + SHO
                                                paragraph.text = "SHO-" + SHO
                                                endGoalieCounter = endGoalieCounter + 1
                                                if endGoalieCounter == 2:
                                                    print "END OF GOAL"
                                                    document.save(newDocDirectory + "/Updated Filev2.docx")
                                                    sys.exit()
                                                flag = 0
                                                print "ROW COUNTER: " + str(row_counter)

                                                break

                                        if paragraph.text.__contains__(name.split(', ')[1]):
                                            print "NAMES2: " + paragraph.text + " " + name.split(', ')[1]
                                            flag = flag + 1
                                            print paragraph.text
                                            print "TEST2"

                                            # scan has a dictionary of all the names found on the website
                                            # the program will check whether or not a first name has already been recorded
                                            # if it has, then the counter will go up from 1 to 2
                                            # if the counter is more than 1, then the program will have to match the specific name
                                            # from the document with a matching number
                                            scan = dic.keys()
                                            for index in scan:
                                                if index.__contains__(name.split(', ')[1]):
                                                    switch = True
                                                    count = count + 1
                                                    dic[name] = count
                                                    count = 1

                                            if switch == True:
                                                continue
                                            dic[name] = count
                            except:
                                print "non men"
                        except IndexError:
                            print "list index out of range"
                except:
                    print "NoneType' object has no attribute 'get"

        print "almost done"
        print "Check on these players' stats."
        for name in dic.keys():
            if dic[name] > 1:
                print name + '\n'
        print " Information might be wrong due to multiple first names appearing."
        document.add_table(1, 2, style=None)
        if gender == "woman":
            document.save(newDocDirectory + "/Women's Soccer Updated File.docx")
        if gender == "man":
            document.save(newDocDirectory + "/Men's Soccer Updated File.docx")
        sys.exit()

    def convertGoalieV2(self, document, tables, goalieName, gender):
        print "GOALIE V2 BEGIN"
        print gender
        endGoalieCounter = 0
        url = ""
        # Reading docx file
        import sys
        from docx import Document
        section = []
        i = 0
        for table in tables:
            section.append(table)
        working = section[1]
        content = []
        flag = 0
        lastNameCounter = 0
        # Import library for reading website
        from bs4 import BeautifulSoup
        import urllib2
        #import pandas as pd
        # download html from link
        #2017 link http://www.dakstats.com/WebSync/Pages/Team/IndividualStats.aspx?association=10&sg=WSO&sea=NAIWSO_2017&team=2409
        if gender == "woman":
            url = "http://www.dakstats.com/WebSync/Pages/Team/IndividualStats.aspx?association=10&sg=WSO&team=2409&sea=NAIWSO_2018"
        if gender == "man":
            url = "http://www.dakstats.com/WebSync/Pages/Team/IndividualStats.aspx?association=10&sg=MSO&team=2623&sea=NAIMSO_2018"
        page = urllib2.urlopen(url)
        soup = BeautifulSoup(page, "html.parser")
        table = soup.find_all("table", {"class": "gridViewReportBuilderWide"})[1]
        avg_scores = []
        names = []
        count = 1
        dic = {}
        row_counter = 0
        for row in table.find_all('tr')[1:]:
            col = row.find_all('td')
            row_counter = row_counter + 1
            # if row_counter > 14:
            #     break
            if len(col) == 13:
                print "SUCCEED13!!!"
                try:
                    name = col[1].find('a', href=True)
                    name = name.get('title')
                    names.append(name.split(', ')[1])
                    if name.split(', ')[1] == goalieName:
                        print "This is it"
                        print name.split(', ')[1]

                        GP = col[2].find(text=True)
                        print "GP: " + GP

                        #
                        # average = field 1
                        GS = col[3].find(text=True)
                        print "GS: " + GS

                        Min = col[4].find(text=True)
                        print "Min: " + Min

                        GA = col[5].find(text=True)
                        print "GA: " + GA

                        GAAvg = col[6].find(text=True)
                        print "GAAvg: " + GAAvg

                        SV = col[7].find(text=True)
                        print "SV: " + SV

                        Pct = col[8].find(text=True)
                        print "Pct: " + Pct

                        W = col[9].find(text=True)
                        print "W: " + W

                        L = col[10].find(text=True)
                        print "L: " + L

                        T = col[11].find(text=True)
                        print "T: " + T

                        SHO = col[12].find(text=True)
                        print "SHO: " + SHO

                        try:
                            print "TESTING"
                            for row in working.rows:
                                for cell in row.cells:
                                    for index, paragraph in enumerate(cell.paragraphs):
                                        content.append(paragraph.text)
                                        switch = False
                                    print str(flag)
                                    print paragraph.text
                                    print name.split(', ')[1]
                                    if flag == 1:
                                        print "W FOUND: " + W
                                        print "W: " + paragraph.text
                                        paragraph.text = W
                                        flag = flag + 1
                                        continue
                                    if flag == 2:
                                        print "L FOUND: " + L
                                        print "L: " + paragraph.text
                                        paragraph.text = L
                                        flag = flag + 1
                                        continue
                                    if flag == 3:
                                        print "T FOUND: " + T
                                        print "T: " + paragraph.text
                                        paragraph.text = T
                                        flag = flag + 1
                                        continue
                                    if flag > 3:
                                        if paragraph.text.__contains__('GP-'):
                                            print 'GP found: ' + GP
                                            paragraph.text = "GP-" + GP

                                        if paragraph.text.__contains__('GS-'):
                                            print 'GS found: ' + GS
                                            paragraph.text = "GS-" + GS

                                        if paragraph.text.__contains__('Min-'):
                                            print 'Min found: ' + Min
                                            paragraph.text = "Min-" + Min

                                        if paragraph.text.__contains__('GA-'):
                                            print 'GA found: ' + GA
                                            paragraph.text = "GA-" + GA

                                        if paragraph.text.__contains__('GAAvg-'):
                                            print 'GAAvg found: ' + GAAvg
                                            paragraph.text = "GAAvg-" + GAAvg

                                        if paragraph.text.__contains__('SV-'):
                                            print 'SV found: ' + SV
                                            paragraph.text = "SV-" + SV

                                        if paragraph.text.__contains__('Pct-'):
                                            print 'Pct found: ' + Pct
                                            paragraph.text = "Pct-" + Pct

                                        if paragraph.text.__contains__('SHO-'):
                                            print 'SHO found: ' + SHO
                                            paragraph.text = "SHO-" + SHO
                                            # endGoalieCounter = endGoalieCounter + 1
                                            # if endGoalieCounter == 2:
                                            #     print "END OF GOAL"
                                            #     document.save(newDocDirectory + "/Updated Filev2.docx")
                                            #     sys.exit()
                                            flag = 0
                                            print "ROW COUNTER: " + str(row_counter)
                                            break

                                    if paragraph.text.__contains__(goalieName):
                                        print "NAMES2: " + paragraph.text + " " + name.split(', ')[1]
                                        flag = flag + 1
                                        print paragraph.text
                                        print "TEST2"

                                        # scan has a dictionary of all the names found on the website
                                        # the program will check whether or not a first name has already been recorded
                                        # if it has, then the counter will go up from 1 to 2
                                        # if the counter is more than 1, then the program will have to match the specific name
                                        # from the document with a matching number
                                        # scan = dic.keys()
                                        # for index in scan:
                                        #     if index.__contains__(name.split(', ')[1]):
                                        #         switch = True
                                        #         count = count + 1
                                        #         dic[name] = count
                                        #         count = 1

                                        if switch == True:
                                            continue
                                        dic[name] = count
                                        # paragraph.text[index + 1] = '.555'
                                        # print paragraph.text[index + 1]
                            try:
                                workingSecondPage = section[3]
                                for row in workingSecondPage.rows:
                                    for cell in row.cells:
                                        for index, paragraph in enumerate(cell.paragraphs):
                                            content.append(paragraph.text)
                                            switch = False
                                        print paragraph.text
                                        print name.split(', ')[1]
                                        if flag == 1:
                                            print "W FOUND: " + W
                                            print "W: " + paragraph.text
                                            paragraph.text = W
                                            flag = flag + 1
                                            continue
                                        if flag == 2:
                                            print "L FOUND: " + L
                                            print "L: " + paragraph.text
                                            paragraph.text = L
                                            flag = flag + 1
                                            continue
                                        if flag == 3:
                                            print "T FOUND: " + T
                                            print "T: " + paragraph.text
                                            paragraph.text = T
                                            flag = flag + 1
                                            continue
                                        if flag > 3:
                                            if paragraph.text.__contains__('GP-'):
                                                print 'GP found: ' + GP
                                                paragraph.text = "GP-" + GP

                                            if paragraph.text.__contains__('GS-'):
                                                print 'GS found: ' + GS
                                                paragraph.text = "GS-" + GS

                                            if paragraph.text.__contains__('Min-'):
                                                print 'Min found: ' + Min
                                                paragraph.text = "Min-" + Min

                                            if paragraph.text.__contains__('GA-'):
                                                print 'GA found: ' + GA
                                                paragraph.text = "GA-" + GA

                                            if paragraph.text.__contains__('GAAvg-'):
                                                print 'GAAvg found: ' + GAAvg
                                                paragraph.text = "GAAvg-" + GAAvg

                                            if paragraph.text.__contains__('SV-'):
                                                print 'SV found: ' + SV
                                                paragraph.text = "SV-" + SV

                                            if paragraph.text.__contains__('Pct-'):
                                                print 'Pct found: ' + Pct
                                                paragraph.text = "Pct-" + Pct

                                            if paragraph.text.__contains__('SHO-'):
                                                print 'SHO found: ' + SHO
                                                paragraph.text = "SHO-" + SHO
                                                # endGoalieCounter = endGoalieCounter + 1
                                                # if endGoalieCounter == 2:
                                                #     print "END OF GOAL"
                                                #     document.save(newDocDirectory + "/Updated Filev2.docx")
                                                #     sys.exit()
                                                flag = 0
                                                # print "ROW COUNTER: " + str(row_counter)

                                                break

                                        if paragraph.text.__contains__(goalieName):
                                            print "NAMES2: " + paragraph.text + " " + name.split(', ')[1]
                                            flag = flag + 1
                                            print paragraph.text
                                            print "TEST2"

                                            # scan has a dictionary of all the names found on the website
                                            # the program will check whether or not a first name has already been recorded
                                            # if it has, then the counter will go up from 1 to 2
                                            # if the counter is more than 1, then the program will have to match the specific name
                                            # from the document with a matching number
                                            # scan = dic.keys()
                                            # for index in scan:
                                            #     if index.__contains__(name.split(', ')[1]):
                                            #         switch = True
                                            #         count = count + 1
                                            #         dic[name] = count
                                            #         count = 1

                                            if switch == True:
                                                continue
                                            dic[name] = count
                            except:
                                print "non men"
                        except IndexError:
                            print "list index out of range"
                except:
                    print "NoneType' object has no attribute 'get"
        print "Function Complete"
        # print "almost done"
        # print "Check on these players' stats."
        # for name in dic.keys():
        #     if dic[name] > 1:
        #         print name + '\n'
        # print " Information might be wrong due to multiple first names appearing."
        # document.add_table(1, 2, style=None)
        # if gender == "woman":
        #     document.save(newDocDirectory + "/Women's Soccer Updated File.docx")
        # if gender == "man":
        #     document.save(newDocDirectory + "/Men's Soccer Updated File.docx")
        # sys.exit()

    def convertWomanSoccer(self):
        gender = "woman"
        import sys
        from docx import Document
        document = Document(oldDocName)
        tables = document.tables
        section = []
        i = 0
        for table in tables:
            section.append(table)
        working = section[1]
        content = []
        flag = 0
        # Import library for reading website
        from bs4 import BeautifulSoup
        import urllib2
        #import pandas as pd
        # download html from link 
        # 2017 http://www.dakstats.com/WebSync/Pages/Team/IndividualStats.aspx?association=10&sg=WSO&sea=NAIWSO_2017&team=2409
        url = "http://www.dakstats.com/WebSync/Pages/Team/IndividualStats.aspx?association=10&sg=WSO&team=2409&sea=NAIWSO_2018"
        page = urllib2.urlopen(url)
        soup = BeautifulSoup(page, "html.parser")
        table = soup.find("table", {"class": "gridViewReportBuilderWide"})
        names = []
        count = 1
        dic = {}
        row_counter = 0
        goalieActivate = 0
        for row in table.find_all('tr')[1:]:
            col = row.find_all('td')
            row_counter = row_counter + 1
            if len(col) == 16:
                print "LENGTH = 16"
                try:
                    name = col[1].find('a', href=True)
                    name = name.get('title')
                    names.append(name.split(', ')[1])
                    if name.split(', ')[1] == "Paula":
                        goalieActivate = 1
                        continue
                    print name.split(', ')[1]

                    GP = col[2].find(text=True)
                    print "GP: " + GP

                    GS = col[3].find(text=True)
                    print "GS: " + GS

                    G = col[4].find(text=True)
                    print "G: " + G

                    A = col[5].find(text=True)
                    print "A: " + A

                    Pts = col[6].find(text=True)
                    print "Pts: " + Pts

                    SH = col[8].find(text=True)
                    print "SH: " + SH

                    SOG = col[9].find(text=True)
                    print "SOG: " + SOG

                    YC = col[11].find(text=True)
                    print "YC: " + YC

                    GWG = col[13].find(text=True)
                    print "GWG: " + GWG

                    PKM = col[14].find(text=True)
                    print "PKM: " + PKM

                    try:
                        print "Begin Convert"
                        for row in working.rows:
                            for cell in row.cells:
                                for index, paragraph in enumerate(cell.paragraphs):
                                    content.append(paragraph.text)
                                    switch = False
                                if flag == 1:
                                    print "G FOUND: " + G
                                    paragraph.text = G
                                    flag = flag + 1
                                    continue
                                if flag == 2:
                                    print "A FOUND: " + A
                                    paragraph.text = A
                                    flag = flag + 1
                                    continue
                                if flag == 3:
                                    print "Pts FOUND: " + Pts
                                    paragraph.text = Pts
                                    flag = flag + 1
                                    continue
                                if flag > 3:
                                    if paragraph.text.__contains__('GP-'):
                                        print 'GP found: ' + GP
                                        paragraph.text = "GP-" + GP

                                    if paragraph.text.__contains__('GS-'):
                                        print 'GS found: ' + GS
                                        paragraph.text = "GS-" + GS

                                    if paragraph.text.__contains__('SH-'):
                                        print 'SH found: ' + SH
                                        paragraph.text = "SH-" + SH

                                    if paragraph.text.__contains__('SOG-'):
                                        print 'SOG found: ' + SOG
                                        paragraph.text = "SOG-" + SOG

                                    if paragraph.text.__contains__('YC-'):
                                        print 'YC found: ' + YC
                                        paragraph.text = "YC-" + YC

                                    if paragraph.text.__contains__('GWG-'):
                                        print 'GWG found: ' + GWG
                                        paragraph.text = "GWG-" + GWG

                                    if paragraph.text.__contains__('PKM-'):
                                        print 'PKM found: ' + PKM
                                        paragraph.text = "PKM-" + PKM
                                        flag = 0
                                        print "ROW COUNTER: " + str(row_counter)
                                        break

                                    # # scan has a dictionary of all the names found on the website
                                    # # the program will check whether or not a first name has already been recorded
                                    # # if it has, then the counter will go up from 1 to 2
                                    # # if the counter is more than 1, then the program will have to match the specific name
                                    # # from the document with a matching number
                                    # scan = dic.keys()
                                    # for index in scan:
                                    #     if index.__contains__(name.split(', ')[1]):
                                    #         switch = True
                                    #         count = count + 1
                                    #         dic[name] = count
                                    #         count = 1
                                    #
                                    # if switch == True:
                                    #     continue
                                    # dic[name] = count

                                if paragraph.text.__contains__(name.split(', ')[1]):
                                    print "NAMES: " + paragraph.text + " " + name.split(', ')[1]
                                    flag = flag + 1
                                    if paragraph.text == "Cheyenne":
                                        print "LOOK HERE NOW"
                                    # self.write(name + '\n')
                                    print paragraph.text
                                    print "TEST2"

                                    # scan has a dictionary of all the names found on the website
                                    # the program will check whether or not a first name has already been recorded
                                    # if it has, then the counter will go up from 1 to 2
                                    # if the counter is more than 1, then the program will have to match the specific name
                                    # from the document with a matching number
                                    scan = dic.keys()
                                    for index in scan:
                                        if index.__contains__(name.split(', ')[1]):
                                            switch = True
                                            count = count + 1
                                            dic[name] = count
                                            count = 1

                                    if switch == True:
                                        continue
                                    dic[name] = count
                    except IndexError:
                        print "list index out of range"
                except:
                    print "NoneType' object has no attribute 'get"

        print "almost done"
        print "Check on these players' stats."
        for name in dic.keys():
            if dic[name] > 1:
                print name + '\n'
        print " Information might be wrong due to multiple first names appearing."
        newDocName = newDocDirectory + "/Women's Soccer Updated File.docx"
        document.save(newDocName)
        if goalieActivate == 1:
            self.convertGoalie(newDocName, gender)
        sys.exit()

    def convertWomanBasketball(self):
        import sys
        from docx import Document
        from docx.shared import Pt
        document = Document(oldDocName)
        tables = document.tables
        #
        style = document.styles['Normal']
        font = style.font
        font.name = 'Calibri (Body)'
        font.size = Pt(8)
        #
        section = []
        i = 0
        for table in tables:
            section.append(table)
        working = section[1]
        content = []
        flag = 0
        # Import library for reading website
        from bs4 import BeautifulSoup
        import urllib2
        #import pandas as pd
        # download html from link
        url = "http://www.dakstats.com/WebSync/Pages/Team/IndividualStats.aspx?association=10&sg=WBB&team=233&sea=NAIWBB_2018"
        page = urllib2.urlopen(url)
        soup = BeautifulSoup(page, "html.parser")
        table = soup.find("table", {"class": "gridViewReportBuilderWide"})
        names = []
        count = 1
        dic = {}
        row_counter = 0
        for row in table.find_all('tr')[1:]:
            col = row.find_all('td')
            row_counter = row_counter + 1
            if len(col) == 25:
                print "LENGTH = 25"
                try:
                    # name*
                    # PCT1*
                    # PCT2*
                    # PCT3*
                    #
                    # GP*
                    # GS*
                    # AVG*
                    # R / G*
                    # A*
                    # BLK*
                    # ST*
                    # P / G
                    name = col[0].find('a', href=True)
                    name = name.get('title')
                    names.append(name.split(', ')[1])
                    print name.split(', ')[1]

                    GP = col[1].find(text=True)
                    print "GP: " + GP

                    GS = col[2].find(text=True)
                    print "GS: " + GS

                    Avg = col[4].find(text=True)
                    print "Avg: " + Avg

                    PCT1 = col[7].find(text=True)
                    print "PCT1: " + PCT1

                    PCT2 = col[10].find(text=True)
                    print "PCT2: " + PCT2

                    PCT3 = col[13].find(text=True)
                    print "PCT3: " + PCT3

                    RG = col[17].find(text=True)
                    print "RG: " + RG

                    A = col[19].find(text=True)
                    print "A: " + A

                    BLK = col[21].find(text=True)
                    print "BLK: " + BLK

                    ST = col[22].find(text=True)
                    print "ST: " + ST

                    PG = col[24].find(text=True)
                    print "PG: " + PG

                    try:
                        print "Begin Convert"
                        for row in working.rows:
                            for cell in row.cells:
                                for index, paragraph in enumerate(cell.paragraphs):
                                    content.append(paragraph.text)
                                    switch = False
                                if flag == 1 and paragraph.text.__contains__("P1:"): 
                                    print "PCT1 FOUND: " + PCT1
                                    paragraph.text = "P1:" + PCT1
                                    flag = flag + 1
                                    print paragraph.text
                                    continue
                                if flag == 2 and paragraph.text.__contains__("P2:"): 
                                    print "PCT2 FOUND: " + PCT2
                                    paragraph.text = "P2:" + PCT2
                                    flag = flag + 1
                                    print paragraph.text
                                    continue
                                if flag == 3 and paragraph.text.__contains__("P3:"): 
                                    print "PCT3 FOUND: " + PCT3
                                    paragraph.text = "P3:" + PCT3
                                    flag = flag + 1
                                    print paragraph.text
                                    continue
                                if flag > 3:
                                    # print paragraph.text
                                    if paragraph.text.__contains__('GP-'):
                                        print 'GP found: ' + GP
                                        paragraph.text = "GP-" + GP
                                        print paragraph.text

                                    if paragraph.text.__contains__('GS-'):
                                        print 'GS found: ' + GS
                                        paragraph.text = "GS-" + GS
                                        print paragraph.text

                                    if paragraph.text.__contains__('AVG-'):
                                        print 'AVG found: ' + Avg
                                        paragraph.text = "AVG-" + Avg

                                    if paragraph.text.__contains__('R/G-'):
                                        print 'RG found: ' + RG
                                        paragraph.text = "RG-" + RG
                                        print paragraph.text

                                    if paragraph.text.__contains__('A-'):
                                        print 'A found: ' + A
                                        paragraph.text = "A-" + A
                                        print paragraph.text

                                    if paragraph.text.__contains__('BLK-'):
                                        print 'BLK found: ' + BLK
                                        paragraph.text = "BLK-" + BLK
                                        print paragraph.text

                                    if paragraph.text.__contains__('ST-'):
                                        print 'ST found: ' + ST
                                        paragraph.text = "ST-" + ST
                                        print paragraph.text

                                    if paragraph.text.__contains__('P/G-'):
                                        print 'PG found: ' + PG
                                        paragraph.text = "P/G-" + PG
                                        flag = 0
                                        print "ROW COUNTER: " + str(row_counter)
                                        break

                                        # # scan has a dictionary of all the names found on the website
                                        # # the program will check whether or not a first name has already been recorded
                                        # # if it has, then the counter will go up from 1 to 2
                                        # # if the counter is more than 1, then the program will have to match the specific name
                                        # # from the document with a matching number
                                        # scan = dic.keys()
                                        # for index in scan:
                                        #     if index.__contains__(name.split(', ')[1]):
                                        #         switch = True
                                        #         count = count + 1
                                        #         dic[name] = count
                                        #         count = 1
                                        #
                                        # if switch == True:
                                        #     continue
                                        # dic[name] = count

                                if paragraph.text == (name.split(', ')[1]):
                                    print "NAMES: " + paragraph.text + " " + name.split(', ')[1]
                                    flag = flag + 1
                                    if paragraph.text == "Cheyenne":
                                        print "LOOK HERE NOW"
                                    # self.write(name + '\n')
                                    print paragraph.text
                                    print "TEST2"

                                    # scan has a dictionary of all the names found on the website
                                    # the program will check whether or not a first name has already been recorded
                                    # if it has, then the counter will go up from 1 to 2
                                    # if the counter is more than 1, then the program will have to match the specific name
                                    # from the document with a matching number
                                    scan = dic.keys()
                                    for index in scan:
                                        if index.__contains__(name.split(', ')[1]):
                                            switch = True
                                            count = count + 1
                                            dic[name] = count
                                            count = 1

                                    if switch == True:
                                        continue
                                    dic[name] = count
                    except IndexError:
                        print "list index out of range"
                except:
                    print "NoneType' object has no attribute 'get"

        print "almost done"
        something = "Check on these players' stats. \n"

        for name in dic.keys():
            if dic[name] > 1:
                name = name + '\n'
                something = something + name
        something = something + "\nInformation might be wrong due to multiple first names appearing."
        import tkMessageBox

        tkMessageBox.showwarning("Warning", something)
        newDocName = newDocDirectory + "/Women's Basketball Updated File.docx"
        document.save(newDocName)
        sys.exit()

    def convertManBasketball(self):
        import sys
        from docx import Document
        from docx.shared import Pt
        document = Document(oldDocName)
        tables = document.tables
        #
        style = document.styles['Normal']
        font = style.font
        font.name = 'Calibri (Body)'
        #
        
        section = []
        i = 0
        for table in tables:
            section.append(table)
        working = section[1]
        content = []
        flag = 0
        # Import library for reading website
        from bs4 import BeautifulSoup
        import urllib2
        #import pandas as pd
        # download html from link
        # http://www.dakstats.com/WebSync/Pages/Team/IndividualStats.aspx?association=10&sg=MBB&sea=NAIMBB_2017&team=556
        url = "http://www.dakstats.com/WebSync/Pages/Team/IndividualStats.aspx?association=10&sg=MBB&team=556&sea=NAIMBB_2018"
        page = urllib2.urlopen(url)
        soup = BeautifulSoup(page, "html.parser")
        table = soup.find("table", {"class": "gridViewReportBuilderWide"})
        names = []
        count = 1
        dic = {}
        row_counter = 0
        for row in table.find_all('tr')[1:]:
            col = row.find_all('td')
            row_counter = row_counter + 1
            if len(col) == 25:
                print "LENGTH = 25"
                try:
                    # name*
                    # PCT1*
                    # PCT2*
                    # PCT3*
                    #
                    # GP*
                    # GS*
                    # AVG*
                    # R / G*
                    # A*
                    # BLK*
                    # ST*
                    # P / G
                    name = col[0].find('a', href=True)
                    name = name.get('title')
                    names.append(name.split(', ')[1])
                    print name.split(', ')[1]

                    GP = col[1].find(text=True)
                    print "GP: " + GP

                    GS = col[2].find(text=True)
                    print "GS: " + GS

                    Avg = col[4].find(text=True)
                    print "Avg: " + Avg

                    PCT1 = col[7].find(text=True)
                    print "PCT1: " + PCT1

                    PCT2 = col[10].find(text=True)
                    print "PCT2: " + PCT2

                    PCT3 = col[13].find(text=True)
                    print "PCT3: " + PCT3

                    RG = col[17].find(text=True)
                    print "RG: " + RG

                    A = col[19].find(text=True)
                    print "A: " + A

                    BLK = col[21].find(text=True)
                    print "BLK: " + BLK

                    ST = col[22].find(text=True)
                    print "ST: " + ST

                    PG = col[24].find(text=True)
                    print "PG: " + PG

                    try:
                        print "Begin Convert"
                        for row in working.rows:
                            for cell in row.cells:
                                for index, paragraph in enumerate(cell.paragraphs):
                                    content.append(paragraph.text)
                                    switch = False
                               
                                if flag == 1 and paragraph.text.__contains__("P1:"): 
                                    # paragraph.style = document.styles['Normal']
                                    font.size = Pt(8)
                                    print "PCT1 FOUND: " + PCT1 + " " + paragraph.text
                                    paragraph.text = "P1:" + PCT1
                                    flag = flag + 1
                                    print paragraph.text
                                    continue
                                if flag == 2 and paragraph.text.__contains__("P2:"):
                                    
                                    print "PCT2 FOUND: " + PCT2 + " " + paragraph.text
                                    paragraph.text = "P2:" + PCT2
                                    flag = flag + 1
                                    print paragraph.text
                                    continue
                                if flag == 3 and paragraph.text.__contains__("P3:"):
                                    
                                    print "PCT3 FOUND: " + PCT3 + " " + paragraph.text
                                    paragraph.text = "P3:" + PCT3
                                    flag = flag + 1
                                    print paragraph.text
                                    continue
                                if flag > 3:
                                    # print paragraph.text
                                    
                                    if paragraph.text.__contains__('GP-'):
                                        print 'GP found: ' + GP
                                        paragraph.text = "GP-" + GP
                                        print paragraph.text

                                    if paragraph.text.__contains__('GS-'):
                                        print 'GS found: ' + GS
                                        paragraph.text = "GS-" + GS
                                        print paragraph.text

                                    if paragraph.text.__contains__('AVG-'):
                                        print 'AVG found: ' + Avg
                                        paragraph.text = "AVG-" + Avg

                                    if paragraph.text.__contains__('R/G-'):
                                        print 'RG found: ' + RG
                                        paragraph.text = "RG-" + RG
                                        print paragraph.text

                                    if paragraph.text.__contains__('A-'):
                                        print 'A found: ' + A
                                        paragraph.text = "A-" + A
                                        print paragraph.text

                                    if paragraph.text.__contains__('BLK-'):
                                        print 'BLK found: ' + BLK
                                        paragraph.text = "BLK-" + BLK
                                        print paragraph.text

                                    if paragraph.text.__contains__('ST-'):
                                        print 'ST found: ' + ST
                                        paragraph.text = "ST-" + ST
                                        print paragraph.text

                                    if paragraph.text.__contains__('P/G-'):
                                        print 'PG found: ' + PG
                                        paragraph.text = "P/G-" + PG
                                        flag = 0
                                        print "ROW COUNTER: " + str(row_counter)
                                        break

                                        # # scan has a dictionary of all the names found on the website
                                        # # the program will check whether or not a first name has already been recorded
                                        # # if it has, then the counter will go up from 1 to 2
                                        # # if the counter is more than 1, then the program will have to match the specific name
                                        # # from the document with a matching number
                                        # scan = dic.keys()
                                        # for index in scan:
                                        #     if index.__contains__(name.split(', ')[1]):
                                        #         switch = True
                                        #         count = count + 1
                                        #         dic[name] = count
                                        #         count = 1
                                        #
                                        # if switch == True:
                                        #     continue
                                        # dic[name] = count
                                # "Website ->" "N'Jai" "Word Doc ->" "N’Jai" 
                                # if paragraph.text == "Corey" and name.split(', ')[1] == "Corey":
                                #     print "Special Instance"
                                #     #print "NAMES: " + paragraph.text + " " + name.split(', ')[1]
                                #     flag = flag + 1
                                #     #print paragraph.text
                                #     #print "TEST2"

                                #     # scan has a dictionary of all the names found on the website
                                #     # the program will check whether or not a first name has already been recorded
                                #     # if it has, then the counter will go up from 1 to 2
                                #     # if the counter is more than 1, then the program will have to match the specific name
                                #     # from the document with a matching number
                                #     scan = dic.keys()
                                #     for index in scan:
                                #         if index.__contains__(name.split(', ')[1]):
                                #             switch = True
                                #             count = count + 1
                                #             dic[name] = count
                                #             count = 1

                                #     if switch == True:
                                #         continue
                                #     dic[name] = count


                                if paragraph.text == (name.split(', ')[1]):
                                    print "NAMES: " + paragraph.text + " " + name.split(', ')[1]
                                    flag = flag + 1
                                    if paragraph.text == "Cheyenne":
                                        print "LOOK HERE NOW"
                                    # self.write(name + '\n')
                                    print paragraph.text
                                    print "TEST2"

                                    # scan has a dictionary of all the names found on the website
                                    # the program will check whether or not a first name has already been recorded
                                    # if it has, then the counter will go up from 1 to 2
                                    # if the counter is more than 1, then the program will have to match the specific name
                                    # from the document with a matching number
                                    scan = dic.keys()
                                    for index in scan:
                                        if index.__contains__(name.split(', ')[1]):
                                            switch = True
                                            count = count + 1
                                            dic[name] = count
                                            count = 1

                                    if switch == True:
                                        continue
                                    dic[name] = count
                    except IndexError:
                        print "list index out of range"
                except:
                    print "NoneType' object has no attribute 'get"

        print "almost done"
        print "Check on these players' stats."
        for name in dic.keys():
            if dic[name] > 1:
                print name + '\n'
        print " Information might be wrong due to multiple first names appearing."
        newDocName = newDocDirectory + "/Men's Basketball Updated File.docx"
        document.save(newDocName)
        sys.exit()


    def convertMenSoccer(self):
        # Reading docx file
        gender = "man"
        import sys
        from docx import Document
        document = Document(oldDocName)
        tables = document.tables
        section = []
        i = 0
        for table in tables:
            section.append(table)
        working = section[1]
        content = []
        flag = 0
        from bs4 import BeautifulSoup
        import urllib2
        #import pandas as pd
        # download html from link
        url = "http://www.dakstats.com/WebSync/Pages/Team/IndividualStats.aspx?association=10&sg=MSO&team=2623&sea=NAIMSO_2018"
        page = urllib2.urlopen(url)
        soup = BeautifulSoup(page, "html.parser")
        table = soup.find("table", {"class": "gridViewReportBuilderWide"})
        names = []
        count = 1
        dic = {}
        row_counter = 0
        goalieActivate = 0
        for row in table.find_all('tr')[1:]:
            col = row.find_all('td')
            row_counter = row_counter + 1
            if len(col) == 16:
                print "LENGTH = 16"
                try:
                    name = col[1].find('a', href=True)
                    name = name.get('title')
                    names.append(name.split(', ')[1])
                    print name.split(', ')[1]
                    if name.split(', ')[1] == "Andres" or name.split(', ')[1] == "Armando" or name.split(', ')[1] == "Juan":
                        print "PLEASE SKIP"
                        goalieActivate = 1
                        goalieName = name.split(', ')[1]
                        print "Goalie Function"
                        self.convertGoalieV2(document, tables, goalieName, gender)
                        continue


                    GP = col[2].find(text=True)
                    print "GP: " + GP

                    GS = col[3].find(text=True)
                    print "GS: " + GS

                    G = col[4].find(text=True)
                    print "G: " + G

                    A = col[5].find(text=True)
                    print "A: " + A

                    Pts = col[6].find(text=True)
                    print "Pts: " + Pts

                    SH = col[8].find(text=True)
                    print "SH: " + SH

                    SOG = col[9].find(text=True)
                    print "SOG: " + SOG

                    YC = col[11].find(text=True)
                    print "YC: " + YC

                    GWG = col[13].find(text=True)
                    print "GWG: " + GWG

                    PKM = col[14].find(text=True)
                    print "PKM: " + PKM

                    try:
                        print "TESTING"
                        for row in working.rows:
                            for cell in row.cells:
                                for index, paragraph in enumerate(cell.paragraphs):
                                    content.append(paragraph.text)
                                    switch = False
                                if flag == 1:
                                    print "G FOUND: " + G
                                    paragraph.text = G
                                    flag = flag + 1
                                    continue
                                if flag == 2:
                                    print "A FOUND: " + A
                                    paragraph.text = A
                                    flag = flag + 1
                                    continue
                                if flag == 3:
                                    print "Pts FOUND: " + Pts
                                    paragraph.text = Pts
                                    flag = flag + 1
                                    continue
                                if flag > 3:
                                    if paragraph.text.__contains__('GP-'):
                                        print 'GP found: ' + GP
                                        paragraph.text = "GP-" + GP

                                    if paragraph.text.__contains__('GS-'):
                                        print 'GS found: ' + GS
                                        paragraph.text = "GS-" + GS

                                    if paragraph.text.__contains__('SH-'):
                                        print 'SH found: ' + SH
                                        paragraph.text = "SH-" + SH

                                    if paragraph.text.__contains__('SOG-'):
                                        print 'SOG found: ' + SOG
                                        paragraph.text = "SOG-" + SOG

                                    if paragraph.text.__contains__('YC-'):
                                        print 'YC found: ' + YC
                                        paragraph.text = "YC-" + YC

                                    if paragraph.text.__contains__('GWG-'):
                                        print 'GWG found: ' + GWG
                                        paragraph.text = "GWG-" + GWG

                                    if paragraph.text.__contains__('PKM-'):
                                        print 'PKM found: ' + PKM
                                        paragraph.text = "PKM-" + PKM
                                        flag = 0
                                        print "ROW COUNTER: " + str(row_counter)
                                        break

                                    # # scan has a dictionary of all the names found on the website
                                    # # the program will check whether or not a first name has already been recorded
                                    # # if it has, then the counter will go up from 1 to 2
                                    # # if the counter is more than 1, then the program will have to match the specific name
                                    # # from the document with a matching number
                                    # scan = dic.keys()
                                    # for index in scan:
                                    #     if index.__contains__(name.split(', ')[1]):
                                    #         switch = True
                                    #         count = count + 1
                                    #         dic[name] = count
                                    #         count = 1
                                    #
                                    # if switch == True:
                                    #     continue
                                    # dic[name] = count

                                if paragraph.text == name.split(' ')[1]:
                                    
                                    print "NAMES: " + paragraph.text + " " + name.split(', ')[1]
                                    flag = flag + 1
                                    if paragraph.text == "Cheyenne":
                                        print "LOOK HERE NOW"
                                    # self.write(name + '\n')
                                    print paragraph.text
                                    print "TEST2"

                                    # scan has a dictionary of all the names found on the website
                                    # the program will check whether or not a first name has already been recorded
                                    # if it has, then the counter will go up from 1 to 2
                                    # if the counter is more than 1, then the program will have to match the specific name
                                    # from the document with a matching number
                                    scan = dic.keys()
                                    for index in scan:
                                        if index.__contains__(name.split(', ')[1]):
                                            switch = True
                                            count = count + 1
                                            dic[name] = count
                                            count = 1

                                    if switch == True:
                                        continue
                                    dic[name] = count
                        try:
                            # flag = 0
                            
                            print "Next Page Starting + " + str(flag)
                            flag = 0
                            print "Next Page Starting + " + str(flag)
                            workingSecondPage = section[3]
                            for row in workingSecondPage.rows:
                                for cell in row.cells:
                                    for index, paragraph in enumerate(cell.paragraphs):
                                        content.append(paragraph.text)
                                        switch = False
                                    if flag == 1:
                                        print "G FOUND: " + G + " true: " + paragraph.text
                                        paragraph.text = G
                                        flag = flag + 1
                                        continue
                                    if flag == 2:
                                        print "A FOUND: " + A + " true: " + paragraph.text
                                        paragraph.text = A
                                        flag = flag + 1
                                        continue
                                    if flag == 3:
                                        print "Pts FOUND: " + Pts + " true: " + paragraph.text
                                        paragraph.text = Pts
                                        flag = flag + 1
                                        continue
                                    if flag > 3:
                                        print "true: " + paragraph.text
                                        if paragraph.text.__contains__('GP-'):
                                            print 'GP found: ' + GP
                                            paragraph.text = "GP-" + GP

                                        if paragraph.text.__contains__('GS-'):
                                            print 'GS found: ' + GS
                                            paragraph.text = "GS-" + GS

                                        if paragraph.text.__contains__('SH-'):
                                            print 'SH found: ' + SH
                                            paragraph.text = "SH-" + SH

                                        if paragraph.text.__contains__('SOG-'):
                                            print 'SOG found: ' + SOG
                                            paragraph.text = "SOG-" + SOG

                                        if paragraph.text.__contains__('YC-'):
                                            print 'YC found: ' + YC
                                            paragraph.text = "YC-" + YC

                                        if paragraph.text.__contains__('GWG-'):
                                            print 'GWG found: ' + GWG
                                            paragraph.text = "GWG-" + GWG

                                        if paragraph.text.__contains__('PKM-'):
                                            print 'PKM found: ' + PKM
                                            paragraph.text = "PKM-" + PKM
                                            flag = 0
                                            print "ROW COUNTER: " + str(row_counter)
                                            # document.save(newDocDirectory + "/Updated File.docx")
                                            break
                                        # scan = dic.keys()
                                        # for index in scan:
                                        #     if index.__contains__(name.split(', ')[1]):
                                        #         switch = True
                                        #         count = count + 1
                                        #         dic[name] = count
                                        #         count = 1
                                        #
                                        # if switch == True:
                                        #     continue
                                        # dic[name] = count

                                    if paragraph.text == (name.split(', ')[1]):
                                        print "NAMES: " + paragraph.text + " " + name.split(', ')[1]
                                        flag = flag + 1
                                        if paragraph.text == "Cheyenne":
                                            print "LOOK HERE NOW"
                                        if paragraph.text == "David":
                                            print "LOOK HERE NOW2"
                                            flag = flag + 1
                                            goalieActivate = 1
                                        # self.write(name + '\n')
                                        print paragraph.text
                                        print "TEST2"

                                        # scan has a dictionary of all the names found on the website
                                        # the program will check whether or not a first name has already been recorded
                                        # if it has, then the counter will go up from 1 to 2
                                        # if the counter is more than 1, then the program will have to match the specific name
                                        # from the document with a matching number
                                        scan = dic.keys()
                                        for index in scan:
                                            if index.__contains__(name.split(', ')[1]):
                                                switch = True
                                                count = count + 1
                                                dic[name] = count
                                                count = 1

                                        if switch == True:
                                            continue
                                        dic[name] = count
                        except:
                            print "No 2nd page"
                    except IndexError:
                        print "list index out of range"
                except:
                    print "NoneType' object has no attribute 'get"

        print "Check on these players' stats."
        for name in dic.keys():
            if dic[name] > 1:
                print name + '\n'
        print " Information might be wrong due to multiple first names appearing."
        newDocName = newDocDirectory + "/Men's Soccer Updated File.docx"
        document.save(newDocName)
        # if goalieActivate == 1:
        # self.convertGoalie(newDocName, gender)
        sys.exit()

    def convertVolleyball(self):
        # Reading docx file
        import sys
        from docx import Document
        document = Document(oldDocName)
        tables = document.tables
        section = []
        i = 0
        for table in tables:
            section.append(table)
        working = section[1]
        content = []
        flag = 0
        # Import library for reading website
        from bs4 import BeautifulSoup
        import urllib2
        #import pandas as pd
        # download html from link
        # 2017: http://www.dakstats.com/WebSync/Pages/Team/IndividualStats.aspx?association=10&sg=WVB&sea=NAIWVB_2017&team=2030

        url = "http://www.dakstats.com/WebSync/Pages/Team/IndividualStats.aspx?association=10&sg=WVB&team=2030&sea=NAIWVB_2018"
        page = urllib2.urlopen(url)
        soup = BeautifulSoup(page, "html.parser")
        table = soup.find("table", {"class": "gridViewReportBuilderWide"})
        names = []
        count = 1
        dic = {}
        row_counter = 0
        for row in table.find_all('tr')[1:]:
            col = row.find_all('td')
            row_counter = row_counter + 1
            if row_counter > 14:
                break
            if len(col) == 30:
                name = col[0].find('a', href=True)
                name = name.get('title')
                names.append(name.split(', ')[1])
                print name.split(', ')[1]
                # if name.split(', ')[1] == "Brooke" or name.split(', ')[1] == 'Morgann':
                #     continue
                if name.split(', ')[1] == "Alyssa":
                    print "ENDING PROGRAM"
                    document.save(newDocDirectory + "/Volleyball New Updated File2.docx")
                    sys.exit()


                GP = col[1].find(text=True)
                print "GP: " + GP

                MP = col[2].find(text=True)
                print "MP: " + MP

                K = col[3].find(text=True)
                print "K: " + K

                A = col[8].find(text=True)
                print "A: " + A

                SA = col[12].find(text=True)
                print "SA: " + SA

                SE = col[14].find(text=True)
                print "SE: " + SE

                DIG = col[20].find(text=True)
                print "DIG: " + DIG

                BS = col[22].find(text=True)
                print "BS: " + BS

                BA = col[23].find(text=True)
                print "BA: " + BA

                TOT = col[24].find(text=True)
                print "TOT: " + TOT

                try:
                    print "Begin Converting"
                    for row in working.rows:
                        for cell in row.cells:
                            for index, paragraph in enumerate(cell.paragraphs):
                                content.append(paragraph.text)
                                switch = False
                            if flag == 1:
                                print "BS FOUND: " + BS
                                paragraph.text = BS
                                flag = flag + 1
                                continue
                            if flag == 2:
                                print "BA FOUND: " + BA
                                paragraph.text = BA
                                flag = flag + 1
                                continue
                            if flag == 3:
                                print "TOT FOUND: " + TOT
                                paragraph.text = TOT
                                flag = flag + 1
                                continue
                            if flag > 3:
                                # print "You are here4"
                                if paragraph.text.__contains__('GP-'):
                                    print 'GP found'
                                    paragraph.text = "GP- " + GP
                                    print paragraph.text

                                if paragraph.text.__contains__('MP-'):
                                    print 'MP found'
                                    paragraph.text = "MP- " + MP
                                    print paragraph.text

                                if paragraph.text.__contains__('K-'):
                                    print 'K found'
                                    paragraph.text = "K- " + K
                                    print paragraph.text

                                if paragraph.text.__contains__('SA-'):
                                    print 'SA found'
                                    paragraph.text = "SA- " + SA
                                    print paragraph.text

                                if paragraph.text.__contains__('SE-'):
                                    print 'SE found'
                                    paragraph.text = "SE- " + SE
                                    print paragraph.text

                                if paragraph.text.__contains__('AST-'):
                                    print 'A found' + A
                                    paragraph.text = "AST- " + A
                                    print paragraph.text

                                if paragraph.text.__contains__('DIG-'):
                                    print 'DIG found'
                                    paragraph.text = "DIG - " + DIG
                                    print paragraph.text
                                    flag = 0
                                    print "ROW COUNTER: " + str(row_counter)
                                    print "__conclude "+ name.split(', ')[1]
                                    # document.save(newDocDirectory + "/Updated File.docx")
                                    if name.split(', ')[1] == "Pamela":
                                        document.save(newDocDirectory + "/Volleyball New Updated File2.docx")
                                        sys.exit()

                                    break
                            # else:
                            #     print "not there yet " + str(flag) 
                                # flag = 0
                                # break

                            if name.split(', ')[1] == 'K.':
                                if paragraph.text == "Ka’imilani":
                                    flag = flag + 1
                                # scan has a dictionary of all the names found on the website
                                # the program will check whether or not a first name has already been recorded
                                # if it has, then the counter will go up from 1 to 2
                                # if the counter is more than 1, then the program will have to match the specific name
                                # from the document with a matching number
                                scan = dic.keys()
                                for index in scan:
                                    if index.__contains__(name.split(', ')[1]):
                                        switch = True
                                        count = count + 1
                                        dic[name] = count
                                        count = 1

                                if switch == True:
                                    continue
                                dic[name] = count

                            # if name.split(', ')[0] == '10 SCOTT':
                            #     print "SCOTT PASSED"
                            #     if paragraph.text == "Jaden":
                            #         flag = flag + 1

                            #     # scan has a dictionary of all the names found on the website
                            #     # the program will check whether or not a first name has already been recorded
                            #     # if it has, then the counter will go up from 1 to 2
                            #     # if the counter is more than 1, then the program will have to match the specific name
                            #     # from the document with a matching number
                            #     scan = dic.keys()
                            #     for index in scan:
                            #         if index.__contains__(name.split(', ')[1]):
                            #             switch = True
                            #             count = count + 1
                            #             dic[name] = count
                            #             count = 1

                            #     if switch == True:
                            #         continue
                            #     dic[name] = count
                            # if paragraph.text == "Pamela":
                            #     print "NOW2"
                            # if name.split(', ')[1] == "Pamela":
                            #     print "NOW3"
                            if name.split(', ')[1] == "Lucyanna":
                                if paragraph.text == "Lucy":
                                    flag = flag + 1
                                scan = dic.keys()
                                for index in scan:
                                    if index.__contains__(name.split(', ')[1]):
                                        switch = True
                                        count = count + 1
                                        dic[name] = count
                                        count = 1

                                if switch == True:
                                    continue
                                dic[name] = count


                            if paragraph.text == name.split(', ')[1]:
                                print "NAMES: " + paragraph.text + " " + name.split(', ')[1]
                                # if name.split(', ')[1] == "Sierra":
                                #     print "NOW"
                                flag = flag + 1
                                # self.write(name + '\n')
                                print paragraph.text
                                # print "TEST2"

                                # scan has a dictionary of all the names found on the website
                                # the program will check whether or not a first name has already been recorded
                                # if it has, then the counter will go up from 1 to 2
                                # if the counter is more than 1, then the program will have to match the specific name
                                # from the document with a matching number
                                scan = dic.keys()
                                for index in scan:
                                    if index.__contains__(name.split(', ')[1]):
                                        switch = True
                                        count = count + 1
                                        dic[name] = count
                                        count = 1

                                if switch == True:
                                    continue
                                dic[name] = count


                except IndexError:
                    print "list index out of range"

        print "almost done"
        # print dic.items()

        print "Check on these players' stats."
        for name in dic.keys():
            if dic[name] > 1:
                print name + '\n'
        print " Information might be wrong due to multiple first names appearing."
        document.add_table(1, 2, style=None)
        document.save(newDocDirectory + "/Volleyball Updated File.docx")
        sys.exit()

if __name__ == '__main__':
    UI().mainloop()